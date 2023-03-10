from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.text import WD_UNDERLINE
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Inches, Cm
from docx.shared import RGBColor
from datetime import datetime
import psycopg2
import os
import sys

user_id = sys.argv[1]

print(user_id)

if len(sys.argv) != 2:
    sys.exit()

# python-docx 패키지 설치
# vscode 터미널에서 pip install python-docx


# docx read
filePath1 = ''
filePath2 = ''
doc = Document(filePath1)
doc1 = Document(filePath2)

# 계약서 글씨체
para = doc.add_paragraph()
run = para.add_run()
run.font.name = '돋움'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '돋움')

# 개인정보동의서 글씨체
para1 = doc1.add_paragraph()
sun = para1.add_run()
sun.font.name = '돋움'
sun._element.rPr.rFonts.set(qn('w:eastAsia'), '돋움')

# docx line number
# for x, paragraph in enumerate(doc.paragraphs):
    # print(str(x) + " : " + paragraph.text)

# for x, paragraph in enumerate(doc1.paragraphs):
#     print(str(x) + " : " + paragraph.text)

# db connect
conn = psycopg2.connect(host='192.168.0.37', dbname='kamco_db1', user='postgres', password='postgres', port=6432)
cursor = conn.cursor()
cursor.execute("SELECT 	user_nm, contr_sido, contr_sig, contr_pnu_cnt, user_addr1, user_addr2, mobile_num, birth, coalesce(priv_chk1, 'N'), coalesce(priv_chk2, 'N'), contr_file, contr_file_path FROM user_info where user_id = '" + user_id + "'")
#cursor.execute("SELECT 	user_nm, contr_sido, contr_sig, contr_pnu_cnt, user_addr1, user_addr2, mobile_num, birth, coalesce(priv_chk1, 'N'), coalesce(priv_chk2, 'N') FROM user_info where user_nm = '신정석'")

row = cursor.fetchall()

if len(row) == 0:
    exit(0)

image_file = ''

# column 추출
for i in row:
    user_nm       = (i[0])
    contr_sido    = (i[1])
    contr_sig     = (i[2])
    contr_pnu_cnt = (i[3])
    user_addr1    = (i[4])
    user_addr2    = (i[5])
    mobile_num    = (i[6])
    birth         = (i[7])
    chk1          = (i[8])
    chk2          = (i[9])
    image_file    = (i[11]) + '/' + (i[10])
    print(image_file)

birthNum = birth[2:8]
birthY   = birth[0:4]
birthM   = birth[4:6]
birthD   = birth[6:8]
now      = datetime.now()
year     = str(now.year)
month    = str(now.month)
day      = str(now.day)

####################################################계약서 수정#########################################################

for p in doc.paragraphs:
    if "관련하여 [ ]" in p.text: 
        p.text = p.text.replace("관련하여 [ ]", "관련하여 [" + user_nm + "]") 
    elif " [ ](시/도) " in p.text:
        p.text = p.text.replace(" [ ](시/도) ", "[ " + contr_sido + " ](시/도)")
    
for p in doc.paragraphs:   
    if "[ ](시/군/구) " in p.text:
        p.text = p.text.replace("[ ](시/군/구) ", "[ " + contr_sig + " ](시/군/구)")
    elif "2023 년     월     일" in p.text:
        p.text = p.text.replace("2023 년     월     일", year + " 년    "+   month + "월   " + day + " 일")

# 3번 *문장 글자색, 하이라이트 처리
josa = doc.paragraphs[9].runs[0]
josa.font.color.rgb = RGBColor(255, 0, 0)
josa.font.highlight_color = WD_COLOR_INDEX.YELLOW
josa.font.name = '돋움'
josa._element.rPr.rFonts.set(qn('w:eastAsia'), '돋움')

# 마지막 장 년 월 일 볼드 처리
yyyymmdd = doc.paragraphs[77].runs[0]
yyyymmdd.bold = True
yyyymmdd.font.name = '돋움'
yyyymmdd._element.rPr.rFonts.set(qn('w:eastAsia'), '돋움')

tables = doc.tables
tables[1].rows[0].cells[4].paragraphs[0].text = user_nm + "    (인)"

# (인) 오른쪽 서명이미지 삽입
pa = tables[1].rows[0].cells[4].paragraphs[0]
run = pa.add_run()
run.add_picture(image_file, width = Inches(0.2), height = Inches(0.15))

tables[1].rows[1].cells[4].paragraphs[0].text = birthNum + " - *******"
tables[1].rows[2].cells[4].paragraphs[0].add_run(user_addr1)
tables[1].rows[3].cells[4].paragraphs[0].add_run(user_addr2)
tables[1].rows[4].cells[4].paragraphs[0].add_run(mobile_num)


####################################################계약서 수정#########################################################


# 테이블 위치 확인
# tables1[0].rows[14].cells[1].paragraphs[0].text = "ssssssssssssssssssssss"

####################################################동의서 수정#########################################################
레
tables1 = doc1.tables
for p in tables1[0].rows[5].cells[1].paragraphs:
    if chk1 == 'Y':
        p.text = p.text.replace("선택적 정보 (동의함 ▢ 동의하지 않음 ▢ )", "선택적 정보 (동의함  동의하지 않음 ▢ )")
    elif chk1 == 'N':
        p.text = p.text.replace("선택적 정보 (동의함 ▢ 동의하지 않음 ▢ )", "선택적 정보 (동의함 ▢ 동의하지 않음  )")
    
for p in tables1[0].rows[13].cells[1].paragraphs:
    if chk2 == 'Y':
        p.text = p.text.replace("선택적 정보 (동의함 ▢ 동의하지 않음 ▢ )", "선택적 정보 (동의함  동의하지 않음 ▢ )")
    elif chk2 == 'N':
        p.text = p.text.replace("선택적 정보 (동의함 ▢ 동의하지 않음 ▢ )", "선택적 정보 (동의함 ▢ 동의하지 않음  )")
                
for p in tables1[1].rows[0].cells[0].paragraphs:
    if "2023년 0월 0일" in p.text:
        p.text = p.text.replace("2023년 0월 0일", year + " 년    "+   month + "월   " + day + " 일")
    elif "서명 또는" in p.text:
        p.text = p.text.replace("서명 또는", user_nm + "     서명 또는")   
    elif "생년월일 :" in p.text:
        p.text = p.text.replace("생년월일 :", "생년월일 :  " + birthY + "." + birthM + "." + birthD)

pa1 = tables1[1].rows[0].cells[0].paragraphs[14]
run1 = pa1.add_run()
run1.add_picture(image_file, width = Inches(0.2), height = Inches(0.15))

for p in doc1.paragraphs:
    if "2023 년    월    일" in p.text:
        p.text = p.text.replace("2023 년    월    일", year + " 년    "+   month + "월   " + day + " 일")
    elif "서명 또는" in p.text:
        p.text = p.text.replace("서명 또는", user_nm + "      서명 또는")   
    elif "생년월일 :                          " in p.text:
        p.text = p.text.replace("생년월일 :                          ", "생년월일 :  " + birthY + "." + birthM + "." + birthD)

nameIn = doc1.paragraphs[6].runs[0]
nameIn.add_picture(image_file, width = Inches(0.2), height = Inches(0.15))


####################################################동의서 수정#########################################################

# 수정 된 docx 저장
doc.save(user_id + "_1.docx")
doc1.save(user_id + "_2.docx")

# vscode 터미널에서 sudo apt-get install unoconv 입력
# 우분투 명령어로 docx -> pdf 파일 추가 생성
os.system('doc2pdf ' + user_id + '_1.docx')
os.system('doc2pdf ' + user_id + '_2.docx')

doc.close()
doc1.close()
conn.close()